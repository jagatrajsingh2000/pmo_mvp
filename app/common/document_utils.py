import io
import json
import re
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple, Union

import aiofiles
from fastapi import UploadFile
from fastapi.responses import StreamingResponse

BASE_DIR = Path(__file__).resolve().parents[2]
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


def sanitize_filename(filename: str, default: str = "report.docx") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", filename or "").strip("-")
    return cleaned or default


async def save_upload_file(upload_file: UploadFile, destination: Union[str, Path]) -> Path:
    dest = Path(destination)
    dest.parent.mkdir(parents=True, exist_ok=True)
    async with aiofiles.open(dest, "wb") as out_file:
        content = await upload_file.read()
        await out_file.write(content)
    return dest


async def save_agent_upload(upload_file: UploadFile, agent_slug: str) -> Tuple[str, Path]:
    file_id = str(uuid.uuid4())
    safe_name = sanitize_filename(upload_file.filename or "upload.bin", "upload.bin")
    destination = UPLOAD_DIR / f"{agent_slug}-{file_id}-{safe_name}"
    await save_upload_file(upload_file, destination)
    return file_id, destination


def extract_text_from_file(path: Union[str, Path]) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_text_docx(path)
    if suffix in (".xlsx", ".xls"):
        return _extract_text_excel(path)
    if suffix == ".csv":
        return _extract_text_csv(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_text_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            continue
    combined = "\n".join(texts)
    if len(combined.strip()) >= 100:
        return combined

    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(str(path))
        ocr_texts = [pytesseract.image_to_string(image) for image in images]
        ocr_combined = "\n".join(ocr_texts)
        if len(ocr_combined.strip()) > len(combined.strip()):
            return ocr_combined
    except Exception:
        pass
    return combined


def _extract_text_docx(path: Path) -> str:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(path))
    parts = []
    table_index = 0

    for block in doc.element.body.iterchildren():
        if block.tag.endswith("}p"):
            paragraph = Paragraph(block, doc)
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                parts.append(f"\nSection: {text}")
            else:
                parts.append(text)
        elif block.tag.endswith("}tbl"):
            table_index += 1
            parts.append(f"\nTable {table_index}:")
            table = Table(block, doc)
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_excel(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(str(path), sheet_name=None, engine="openpyxl")
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        parts.append(df.to_csv(index=False))
    return "\n".join(parts)


def _extract_text_csv(path: Path) -> str:
    import pandas as pd

    return pd.read_csv(str(path)).to_csv(index=False)


def normalize_report_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    if "resolved" in payload and isinstance(payload.get("resolved"), dict):
        return payload["resolved"]
    if "generated" in payload and isinstance(payload.get("generated"), dict):
        return payload["generated"]
    return payload


def write_json_output(file_id: str, agent_slug: str, payload: Dict[str, Any]) -> Path:
    output_path = OUTPUT_DIR / f"{agent_slug}-{file_id}.json"
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return output_path


def docx_response(title: str, payload: Dict[str, Any], filename: str) -> StreamingResponse:
    document = build_docx_document(title, payload)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    safe_name = sanitize_filename(filename, "agent-report.docx")
    if not safe_name.lower().endswith(".docx"):
        safe_name = f"{safe_name}.docx"
    return StreamingResponse(
        buffer,
        media_type=DOCX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


def build_docx_document(title: str, payload: Dict[str, Any]):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)

    heading = document.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(32, 99, 90)

    _render_value(document, payload, 1)
    return document


def _render_value(document, value: Any, level: int, key: str = "") -> None:
    if isinstance(value, dict):
        if key:
            document.add_heading(_title(key), min(level, 3))
        scalar_rows = [(k, v) for k, v in value.items() if _is_scalar(v)]
        nested_rows = [(k, v) for k, v in value.items() if not _is_scalar(v)]
        if scalar_rows:
            table = document.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Field"
            table.rows[0].cells[1].text = "Value"
            for row_key, row_value in scalar_rows:
                cells = table.add_row().cells
                cells[0].text = _title(row_key)
                cells[1].text = _stringify(row_value)
        for nested_key, nested_value in nested_rows:
            _render_value(document, nested_value, level + 1, nested_key)
        return

    if isinstance(value, list):
        if key:
            document.add_heading(_title(key), min(level, 3))
        if not value:
            document.add_paragraph(
                "Not provided in source. Impact: confirm whether this area is intentionally out of scope or requires follow-up input."
            )
            return
        if all(isinstance(item, dict) for item in value):
            _render_dict_table(document, value)
            return
        for item in value:
            document.add_paragraph(_stringify(item), style="List Bullet")
        return

    if key:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{_title(key)}: ").bold = True
        paragraph.add_run(_stringify(value))
    else:
        document.add_paragraph(_stringify(value))


def _render_dict_table(document, rows: List[Dict[str, Any]]) -> None:
    columns = _collect_columns(rows)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = _title(column)
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = _stringify(row.get(column))


def _collect_columns(rows: Iterable[Dict[str, Any]]) -> List[str]:
    columns = []
    for row in rows:
        for key, value in row.items():
            if key not in columns and _is_scalar(value):
                columns.append(key)
        for key, value in row.items():
            if key not in columns and not _is_scalar(value):
                columns.append(key)
    return columns or ["value"]


def _is_scalar(value: Any) -> bool:
    return value is None or isinstance(value, (str, int, float, bool))


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _title(value: str) -> str:
    return re.sub(r"[_-]+", " ", str(value)).strip().title()
